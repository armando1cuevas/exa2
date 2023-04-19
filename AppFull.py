import streamlit as st
import openai
import json
import chardet
import os
import base64
import PyPDF2
from docx import Document
from io import BytesIO
from docx import Document

openai.api_key = st.secrets["OPENAI_KEY"]
os.environ["OPENAI_API_KEY"] = openai.api_key

def generate_questions_gpt35_turbo(text, num_questions, question_type, num_options=4):



    if question_type == "Opción múltiple":
        gpt_prompt = f"Por favor, genera {num_questions} preguntas de opción múltiple con {num_options} opciones de respuesta de las cuales solo una es correcta. Agrega un asterisco al final de la respuesta correcta."
        role_content = "You are a helpful assistant that generates multiple choice questions based on the provided text."
    elif question_type == "Falso/Verdadero":
        gpt_prompt = f"Por favor, genera {num_questions} preguntas de tipo Falso o Verdadero. Incluye las posibles respuestas A) Verdadero o B) Falso, agregando un asterisco al final de la respuesta correcta."
        role_content = "You are a helpful assistant that generates true or false questions based on the provided text."
    else:
        gpt_prompt = f"Por favor, genera {num_questions} preguntas del tipo fill-in-the-blank (completar el espacio en blanco) con 4 opciones de respuesta cada una. El enunciado de la pregunta debe mostrar una línea punteada '_____' en lugar de la parte omitida de la oración. Agrega un asterisco al final de la respuesta correcta."
        role_content = "You are a helpful assistant that generates fill-in-the-blank questions based on the provided text."

    prompt = [
        {"role": "system", "content": role_content},
        {"role": "user", "content": text},
        {"role": "assistant", "content": gpt_prompt},
    ]

    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=prompt,
        temperature=0.5,
        max_tokens=2000,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
    )

    return response['choices'][0]['message']['content'].strip()

def read_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in range(len(pdf_reader.pages)):
        text += pdf_reader.pages[page].extract_text()
    return text

def read_docx(file):
    doc = Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def text_to_docx(questions):
    doc = Document()
    doc.add_paragraph(questions)
    byte_io = BytesIO()
    doc.save(byte_io)
    return byte_io.getvalue()

def get_binary_file_downloader_html(bin_file, file_label='File'):
    with st.spinner('Generando archivo...'):
        b64 = base64.b64encode(bin_file).decode()
        href = f'<a href="data:application/octet-stream;base64,{b64}" download="{file_label}.docx">Descargar archivo .docx</a>'
    return href

st.title("Generador de preguntas GPT-4")
st.write("Sube un archivo de texto y genera preguntas de Opción múltiple, Falso/Verdadero o Fill-in-the-blank basadas en su contenido.")

uploaded_file = st.file_uploader("Sube un archivo de texto", type=['txt', 'pdf', 'docx'])

if uploaded_file:
    file_extension = uploaded_file.name.split(".")[-1].lower()

    if file_extension == "txt":
        raw_data = uploaded_file.read()
        detected_encoding = chardet.detect(raw_data)
        text = raw_data.decode(detected_encoding["encoding"])
    elif file_extension == "pdf":
        text = read_pdf(uploaded_file)
    elif file_extension == "docx":
        text = read_docx(uploaded_file)
    else:
        st.error("Formato de archivo no soportado. Por favor, sube un archivo TXT, PDF o DOCX.")

    question_type = st.selectbox("Selecciona el tipo de preguntas a generar", ["Opción múltiple", "Falso/Verdadero", "Fill-in-the-blank"])

    if question_type == "Opción múltiple":
        num_options = st.number_input("Selecciona el número de posibles respuestas para cada pregunta:", min_value=2, max_value=10, value=4, step=1)
    else:
        num_options = 4

    num_questions = st.number_input("Número de preguntas a generar", min_value=1, value=5)

    if st.button("Generar preguntas"):
        questions = generate_questions_gpt35_turbo(text, num_questions, question_type, num_options)
        st.write(questions.replace("*", "\\*"))

        docx_file = text_to_docx(questions)
        st.markdown(get_binary_file_downloader_html(docx_file, 'preguntas_y_respuestas'), unsafe_allow_html=True)


